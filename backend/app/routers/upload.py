from fastapi import APIRouter, UploadFile, File, Form, Depends, HTTPException
from sqlalchemy.orm import Session
import os
import shutil
import hashlib
import re
from datetime import datetime

from app.database import get_db, Document

router = APIRouter()

UPLOAD_DIR = "backend/uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)


def extract_text_from_file(file_path: str, filename: str) -> str:
    """Извлечение текста из файла"""
    text = ""
    try:
        if filename.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        elif filename.endswith('.pdf'):
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            for page in reader.pages:
                text += page.extract_text() or ""
        elif filename.endswith(('.docx', '.doc')):
            from docx import Document
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        elif filename.endswith('.rtf'):
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
                text = re.sub(r'\\[a-z]+', ' ', text)
                text = re.sub(r'\{.*?\}', '', text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка извлечения текста: {str(e)}")

    text = re.sub(r'\s+', ' ', text).strip()
    return text


@router.post("/file")
async def upload_file(
        file: UploadFile = File(...),
        file_type: str = Form("other"),
        db: Session = Depends(get_db)
):
    """Загрузить файл для анализа"""
    try:
        allowed_extensions = ['.txt', '.pdf', '.docx', '.doc', '.rtf']
        file_ext = os.path.splitext(file.filename)[1].lower()

        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"Неподдерживаемый формат. Поддерживаются: {', '.join(allowed_extensions)}"
            )

        file_path = os.path.join(UPLOAD_DIR, file.filename)
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        text = extract_text_from_file(file_path, file.filename)

        if not text or len(text) < 10:
            raise HTTPException(
                status_code=400,
                detail="Текст слишком короткий или не удалось извлечь содержимое"
            )

        content_hash = hashlib.sha256(text.encode('utf-8')).hexdigest()

        existing_doc = db.query(Document).filter(
            Document.content_hash == content_hash
        ).first()

        if existing_doc:
            return {
                "success": True,
                "document_id": existing_doc.id,
                "filename": existing_doc.filename,
                "file_type": existing_doc.file_type,
                "word_count": existing_doc.word_count,
                "file_size": existing_doc.file_size,
                "uploaded_at": existing_doc.uploaded_at.isoformat(),
                "is_duplicate": True,
                "message": "Файл с таким содержимым уже загружен"
            }

        doc = Document(
            filename=file.filename,
            content=text,
            content_hash=content_hash,
            source="user_upload",
            file_type=file_type,
            file_size=os.path.getsize(file_path),
            word_count=len(text.split())
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)

        return {
            "success": True,
            "document_id": doc.id,
            "filename": doc.filename,
            "file_type": doc.file_type,
            "word_count": doc.word_count,
            "file_size": doc.file_size,
            "uploaded_at": doc.uploaded_at.isoformat(),
            "is_duplicate": False,
            "message": "Файл успешно загружен"
        }

    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))